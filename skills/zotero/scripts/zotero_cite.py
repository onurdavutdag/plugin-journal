#!/usr/bin/env python3
"""Word citation/bibliography engine — mimics the Zotero Word plugin buttons.

Part of the `zotero` skill. Reads the user's real Zotero library via
`zotero_lib.py` (same directory) and processes citation markers in a .docx.

Marker syntax inside the document (either form, one or more keys, ';'-separated):
    {{zref:ITEMKEY}}        {{zref:KEY1;KEY2}}
    [@ITEMKEY]              [@KEY1;KEY2]
ITEMKEY is the 8-char Zotero item key (e.g. F5RI4K5K).

Modes:
  --mode field (default) — every marker becomes a REAL Zotero Word-integration
      field (ADDIN ZOTERO_ITEM CSL_CITATION + ZOTERO_PREF + ZOTERO_BIBL), so the
      user's actual Zotero app recognizes, refreshes and restyles the citations
      from Word (verified against the user's Zotero 7 + Word). Re-running only
      converts NEW markers; existing ZOTERO_* fields belong to the Zotero app
      and are never touched.
  --mode text — legacy static text (WJ-wrapped), refreshed only by this script.

Buttons implemented:
  refresh (default) — Add/Edit Citation + Refresh: replace every marker with the
      in-text citation in the chosen style, (re)number in order of appearance,
      and (re)write the bibliography at the end. Text mode is idempotent within
      the script; field mode hands renumbering to the Zotero app.
  unlink — freeze (text mode): keep the rendered citation text, remove marker
      bookkeeping so the document can no longer be refreshed. For field mode use
      Zotero's own Unlink Citations button instead.

Styles:
  --style vancouver  (default) numeric [1], bibliography in citation order.
      Field mode maps to CSL style id http://www.zotero.org/styles/vancouver.
  --style author-date          (Author, Year), bibliography alphabetical.
      Field mode maps to CSL style id http://www.zotero.org/styles/apa.
  Journal-specific styling beyond these two is handled WITHIN the zotero skill
  (read CSL rules, apply here) — or, in field mode, directly from the Zotero
  app (Document Preferences). Docx citation/bibliography authority lives only
  in this skill; no other skill/agent formats or edits the reference list.

Red-revision rule (global CLAUDE.md): when updating an EXISTING document
(default), all text this script inserts is red (RGB 255,0,0). Pass --no-red for
brand-new documents built from scratch.

Output: JSON report to stdout {processed_markers, unknown_keys, bibliography_count, output}.

Usage:
    python zotero_cite.py --docx makale.docx                       # refresh, Vancouver
    python zotero_cite.py --docx makale.docx --style author-date
    python zotero_cite.py --docx makale.docx --action unlink
    python zotero_cite.py --docx makale.docx --out makale_atifli.docx --no-red
"""
import argparse
import json
import os
import random
import re
import subprocess
import sys

try:
    import docx  # python-docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
except ImportError:
    print(json.dumps({"error": "no_python_docx",
                      "message": "python-docx kurulu değil: pip install python-docx"}))
    sys.exit(0)

RED = RGBColor(0xFF, 0x00, 0x00)
MARKER_RE = re.compile(r"\{\{zref:([A-Z0-9;\s]+)\}\}|\[@([A-Z0-9;\s]+)\]")
BIB_HEADINGS = ("Kaynaklar", "Kaynakça", "References", "Bibliography")
# Invisible bookmark char pair used to keep refresh idempotent: rendered
# citations are wrapped as ⁠{{zref:...}}⁠<visible text>⁠ in a
# single run's text? Too fragile across runs — instead we KEEP the marker in
# the document and render the citation right after it, wrapped in U+2060
# word-joiners, and strip that rendered part on each refresh.
WJ = "⁠"  # word joiner, invisible in Word

_RENDERED_RE = re.compile(WJ + r"[^" + WJ + r"]*" + WJ)


# ------------------------------------------------------------ library -------

def load_library():
    lib_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "zotero_lib.py")
    out = subprocess.run([sys.executable, lib_script, "--items"],
                         capture_output=True, text=True, encoding="utf-8")
    data = json.loads(out.stdout)
    if isinstance(data, dict) and data.get("error"):
        raise SystemExit(json.dumps(data, ensure_ascii=False))
    return {it["key"]: it for it in data}


def load_account():
    """Account identifiers from zotero_lib --status (for field URIs)."""
    lib_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "zotero_lib.py")
    out = subprocess.run([sys.executable, lib_script, "--status"],
                         capture_output=True, text=True, encoding="utf-8")
    try:
        return json.loads(out.stdout).get("account") or {}
    except (json.JSONDecodeError, AttributeError):
        return {}


# ------------------------------------------------------- Zotero fields ------
# Real Word-integration fields (ADDIN ZOTERO_*) the actual Zotero app manages.
# Structure verified against the user's Zotero 7 + Word (feasibility spike).

STYLE_IDS = {
    "vancouver": "http://www.zotero.org/styles/vancouver",
    "author-date": "http://www.zotero.org/styles/apa",
}

_CSL_TYPES = {
    "journalArticle": "article-journal", "book": "book", "bookSection": "chapter",
    "conferencePaper": "paper-conference", "thesis": "thesis", "report": "report",
    "webpage": "webpage", "preprint": "article", "manuscript": "manuscript",
}

_CITE_INSTR_RE = re.compile(r"ADDIN ZOTERO_ITEM CSL_CITATION\s+(\{.*)", re.S)


def _rand_id(n=8):
    return "".join(random.choice("0123456789ABCDEF") for _ in range(n))


def item_uri(key, account):
    uid = account.get("user_id")
    if uid:
        return f"http://zotero.org/users/{uid}/items/{key}"
    return f"http://zotero.org/users/local/{account.get('local_user_key', 'local')}/items/{key}"


def csl_item_data(it, uri):
    """zotero_lib record -> CSL-JSON itemData (what Zotero embeds in the field)."""
    d = {
        "id": uri,
        "type": _CSL_TYPES.get(it.get("itemType"), "document"),
        "title": it.get("title"),
        "container-title": it.get("container-title"),
        "journalAbbreviation": it.get("journalAbbreviation"),
        "volume": it.get("volume"),
        "issue": it.get("issue"),
        "page": it.get("pages"),
        "DOI": it.get("DOI"),
        "URL": it.get("url"),
    }
    auths = [{"family": c.get("family", ""), "given": c.get("given", "")}
             for c in it.get("creators", []) if c.get("type") in (None, "author")]
    if auths:
        d["author"] = auths
    if it.get("year"):
        d["issued"] = {"date-parts": [[int(it["year"])]]}
    if it.get("PMID"):
        d["PMID"] = it["PMID"]
    return {k: v for k, v in d.items() if v is not None}


def _add_field(p, instr, result_text, red, breaks=False):
    """Append one complex Word field (begin|instrText|separate|result|end) to p."""
    def raw_run():
        r = OxmlElement("w:r")
        p._p.append(r)
        return r

    def fld(t):
        r = raw_run()
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), t)
        r.append(fc)

    fld("begin")
    r = raw_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " " + instr + " "
    r.append(it)
    fld("separate")
    if result_text:
        lines = result_text.split("\n") if breaks else [result_text]
        run = p.add_run(lines[0])
        if red:
            run.font.color.rgb = RED
        for line in lines[1:]:
            run.add_break()
            run2 = p.add_run(line)
            if red:
                run2.font.color.rgb = RED
    fld("end")


def _scan_field_instrs(doc):
    """Full instruction text of every complex field, in document order."""
    fields = []
    depth = 0
    for el in doc.element.body.iter():
        if el.tag == qn("w:fldChar"):
            t = el.get(qn("w:fldCharType"))
            if t == "begin":
                depth += 1
                if depth == 1:
                    fields.append("")
            elif t == "end" and depth:
                depth -= 1
        elif el.tag == qn("w:instrText") and depth:
            fields[-1] += el.text or ""
    return fields


def _existing_zotero_state(doc):
    """(citation field count, item keys in order, has_pref, has_bibl)."""
    instrs = _scan_field_instrs(doc)
    keys, n_cites = [], 0
    has_pref = any("ADDIN ZOTERO_PREF" in s for s in instrs)
    has_bibl = any("ADDIN ZOTERO_BIBL" in s for s in instrs)
    for s in instrs:
        m = _CITE_INSTR_RE.search(s)
        if not m:
            continue
        n_cites += 1
        try:
            data = json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            continue
        for ci in data.get("citationItems", []):
            for uri in ci.get("uris", []):
                k = uri.rstrip("/").rsplit("/", 1)[-1].upper()
                if k not in keys:
                    keys.append(k)
    return n_cites, keys, has_pref, has_bibl


def _prefs_xml(style):
    return (
        '<data data-version="3" zotero-version="7.0">'
        f'<session id="{_rand_id(8)}"/>'
        f'<style id="{STYLE_IDS[style]}" locale="en-US" '
        'hasBibliography="1" bibliographyStyleHasBeenSet="0"/>'
        '<prefs><pref name="fieldType" value="Field"/>'
        '<pref name="automaticJournalAbbreviations" value="true"/>'
        '<pref name="noteType" value="0"/></prefs></data>'
    )


def _insert_pref_field(doc, style):
    """Document-preferences field in its own paragraph at the top of the body."""
    p = doc.add_paragraph()
    _add_field(p, "ADDIN ZOTERO_PREF_1 " + _prefs_xml(style), "", red=False)
    body = doc.element.body
    body.remove(p._p)
    body.insert(0, p._p)


def citation_field_instr(keys, lib, account, style, order):
    """Build the ADDIN ZOTERO_ITEM instruction + visible result for one marker."""
    citation_items, nums_or_cites = [], []
    for k in keys:
        it = lib[k]
        uri = item_uri(k, account)
        citation_items.append({"id": uri, "uris": [uri], "itemData": csl_item_data(it, uri)})
        if style == "vancouver":
            nums_or_cites.append(str(order.index(k) + 1))
        else:
            nums_or_cites.append(author_date_intext(it)[1:-1])
    if style == "vancouver":
        visible = "[" + ",".join(nums_or_cites) + "]"
    else:
        visible = "(" + "; ".join(nums_or_cites) + ")"
    citation = {
        "citationID": _rand_id(8),
        "properties": {"formattedCitation": visible, "plainCitation": visible, "noteIndex": 0},
        "citationItems": citation_items,
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    return "ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(citation, ensure_ascii=False), visible


def refresh_fields(doc, lib, account, style, red):
    """Convert every marker into a real Zotero citation field.

    Existing ADDIN ZOTERO_* fields are the Zotero app's property — never touched.
    Returns (order, unknown, processed, existing_cites).
    """
    existing_cites, order, has_pref, has_bibl = _existing_zotero_state(doc)
    unknown, processed = [], 0

    for p in _iter_paragraphs(doc):
        raw = _para_text(p)
        if not MARKER_RE.search(raw) and WJ not in raw:
            continue
        text = _strip_rendered(raw)  # drop any legacy text-mode rendering
        matches = list(MARKER_RE.finditer(text))
        if not matches:
            if raw != text:
                _rewrite_paragraph(p, text, [], red)
            continue

        # resolve keys first so `order` is complete before numbering
        per_marker = []
        for m in matches:
            keys = [k.strip().upper() for k in (m.group(1) or m.group(2)).split(";") if k.strip()]
            good = []
            for k in keys:
                if k in lib:
                    good.append(k)
                    if k not in order:
                        order.append(k)
                elif k not in unknown:
                    unknown.append(k)
            per_marker.append(good)

        # rebuild paragraph: plain text segments + citation fields
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        last = 0
        for m, good in zip(matches, per_marker):
            if m.start() > last:
                p.add_run(text[last:m.start()])
            if good:
                instr, visible = citation_field_instr(good, lib, account, style, order)
                _add_field(p, instr, visible, red)
                processed += 1
            else:  # unknown key(s): keep the marker so nothing is silently lost
                p.add_run(text[m.start():m.end()])
            last = m.end()
        if last < len(text):
            p.add_run(text[last:])

    return order, unknown, processed, existing_cites, has_pref, has_bibl


def write_bibliography_field(doc, lib, order, style, red, heading):
    """Heading paragraph + one ADDIN ZOTERO_BIBL field with rendered entries."""
    if not order:
        return 0
    h = doc.add_paragraph()
    run = h.add_run(heading)
    run.bold = True
    if red:
        run.font.color.rgb = RED
    items = [lib[k] for k in order if k in lib]
    if style == "author-date":
        items = sorted(items, key=_sort_key_author_date)
        entries = [author_date_entry(it) for it in items]
    else:
        entries = [f"{i}. {vancouver_entry(it)}" for i, it in enumerate(items, 1)]
    p = doc.add_paragraph()
    _add_field(p, 'ADDIN ZOTERO_BIBL {"uncited":[],"omitted":[],"custom":[]} CSL_BIBLIOGRAPHY',
               "\n".join(entries), red, breaks=True)
    return len(entries)


# ------------------------------------------------------------ formatting ----

def _authors_vancouver(creators):
    """Surname Initials, first 6 then et al. (citation-format.md rule)."""
    auths = [c for c in creators if c.get("type") == "author"] or creators
    names = []
    for c in auths[:6]:
        fam = c.get("family", "").strip()
        giv = c.get("given", "").strip()
        initials = "".join(w[0] for w in re.split(r"[\s\-.]+", giv) if w)
        names.append((fam + " " + initials).strip())
    s = ", ".join(n for n in names if n)
    if len(auths) > 6:
        s += ", et al."
    return s


def vancouver_entry(it):
    """One numbered reference-list entry (without the leading number)."""
    parts = []
    a = _authors_vancouver(it.get("creators", []))
    if a:
        parts.append(a + ".")
    if it.get("title"):
        parts.append(it["title"].rstrip(".") + ".")
    j = it.get("journalAbbreviation") or it.get("container-title")
    if j:
        parts.append(j + ".")
    tail = ""
    if it.get("year"):
        tail = it["year"]
    if it.get("volume"):
        tail += ";" + it["volume"]
        if it.get("issue"):
            tail += "(" + it["issue"] + ")"
    if it.get("pages"):
        tail += ":" + it["pages"]
    if tail:
        parts.append(tail + ".")
    if it.get("DOI"):
        parts.append("doi:" + it["DOI"] + ".")
    if it.get("PMID"):
        parts.append("PMID: " + it["PMID"] + ".")
    return " ".join(parts)


def author_date_intext(it):
    auths = [c for c in it.get("creators", []) if c.get("type") == "author"] or it.get("creators", [])
    if not auths:
        fam = it.get("title", "Anon")[:20]
    elif len(auths) == 1:
        fam = auths[0]["family"]
    elif len(auths) == 2:
        fam = auths[0]["family"] + " & " + auths[1]["family"]
    else:
        fam = auths[0]["family"] + " et al."
    return "(" + fam + ", " + (it.get("year") or "t.y.") + ")"


def author_date_entry(it):
    """APA-flavoured author-date reference entry."""
    auths = [c for c in it.get("creators", []) if c.get("type") == "author"] or it.get("creators", [])
    names = []
    for c in auths:
        giv = c.get("given", "").strip()
        initials = ". ".join(w[0] for w in re.split(r"[\s\-.]+", giv) if w)
        names.append(c.get("family", "") + (", " + initials + "." if initials else ""))
    astr = ", ".join(names[:-1]) + (", & " + names[-1] if len(names) > 1 else (names[0] if names else ""))
    parts = [astr, "(" + (it.get("year") or "t.y.") + ")."]
    if it.get("title"):
        parts.append(it["title"].rstrip(".") + ".")
    j = it.get("container-title") or it.get("journalAbbreviation")
    if j:
        seg = j
        if it.get("volume"):
            seg += ", " + it["volume"]
            if it.get("issue"):
                seg += "(" + it["issue"] + ")"
        if it.get("pages"):
            seg += ", " + it["pages"]
        parts.append(seg + ".")
    if it.get("DOI"):
        parts.append("https://doi.org/" + it["DOI"])
    return " ".join(p for p in parts if p)


def _sort_key_author_date(it):
    cs = it.get("creators", [])
    fam = cs[0]["family"].lower() if cs else ""
    return (fam, it.get("year") or "")


# ------------------------------------------------------------ docx ops ------

def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _para_text(p):
    return "".join(r.text for r in p.runs)


def _strip_rendered(text):
    """Remove previously rendered citation text (idempotent refresh)."""
    return _RENDERED_RE.sub("", text)


def _remove_old_bibliography(doc):
    """Delete a bibliography section previously written by this script."""
    body_paras = doc.paragraphs
    start = None
    for i, p in enumerate(body_paras):
        if p.text.strip().rstrip(WJ).lstrip(WJ) in BIB_HEADINGS and WJ in p.text:
            start = i
            break
    if start is None:
        return
    for p in body_paras[start:]:
        el = p._element
        el.getparent().remove(el)


def _rewrite_paragraph(p, new_text, red_spans, red):
    """Replace paragraph text, coloring the spans listed in red_spans.

    red_spans: list of (start, end) into new_text that must be red (inserted text).
    Existing formatting of the paragraph body is not preserved run-by-run for
    marker paragraphs (acceptable: markers live in author-controlled sentences),
    but the paragraph style stays.
    """
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    pos = 0
    spans = sorted(red_spans)
    for s, e in spans:
        if s > pos:
            p.add_run(new_text[pos:s])
        run = p.add_run(new_text[s:e])
        if red:
            run.font.color.rgb = RED
        pos = e
    if pos < len(new_text):
        p.add_run(new_text[pos:])


# ------------------------------------------------------------ main flow -----

def refresh(doc, lib, style, red):
    """Render/refresh all markers; return (order, unknown, count)."""
    order = []          # item keys in order of first appearance
    unknown = []
    processed = 0

    for p in _iter_paragraphs(doc):
        raw = _para_text(p)
        if WJ not in raw and not MARKER_RE.search(raw):
            continue
        text = _strip_rendered(raw)
        if not MARKER_RE.search(text):
            if raw != text:
                _rewrite_paragraph(p, text, [], red)
            continue

        new_text = ""
        red_spans = []
        last = 0
        for m in MARKER_RE.finditer(text):
            keys = [k.strip().upper() for k in (m.group(1) or m.group(2)).split(";") if k.strip()]
            nums_or_cites = []
            for k in keys:
                it = lib.get(k)
                if it is None:
                    if k not in unknown:
                        unknown.append(k)
                    continue
                if k not in order:
                    order.append(k)
                if style == "vancouver":
                    nums_or_cites.append(str(order.index(k) + 1))
                else:
                    # bare "Author, Year" — merged into one paren below
                    nums_or_cites.append(author_date_intext(it)[1:-1])
            if style == "vancouver":
                cite = "[" + ",".join(nums_or_cites) + "]" if nums_or_cites else "[?]"
            else:
                cite = "(" + "; ".join(nums_or_cites) + ")" if nums_or_cites else "(?)"
            marker = text[m.start():m.end()]
            rendered = marker + WJ + cite + WJ
            new_text += text[last:m.start()] + rendered
            # only the visible citation part is "inserted text" -> red
            cite_start = len(new_text) - len(cite) - 1
            red_spans.append((cite_start, cite_start + len(cite)))
            last = m.end()
            processed += 1
        new_text += text[last:]
        _rewrite_paragraph(p, new_text, red_spans, red)

    return order, unknown, processed


def write_bibliography(doc, lib, order, style, red, heading):
    _remove_old_bibliography(doc)
    if not order:
        return 0
    h = doc.add_paragraph()
    run = h.add_run(heading + WJ)  # WJ tags it as ours for future removal
    run.bold = True
    if red:
        run.font.color.rgb = RED
    items = [lib[k] for k in order]
    if style == "author-date":
        items = sorted(items, key=_sort_key_author_date)
        entries = [author_date_entry(it) for it in items]
    else:
        entries = [f"{i}. {vancouver_entry(it)}" for i, it in enumerate(items, 1)]
    for e in entries:
        p = doc.add_paragraph()
        run = p.add_run(e)
        if red:
            run.font.color.rgb = RED
    return len(entries)


def unlink(doc):
    """Freeze rendered citations: drop markers + WJ wrappers, keep visible text."""
    count = 0
    for p in _iter_paragraphs(doc):
        raw = _para_text(p)
        if WJ not in raw and not MARKER_RE.search(raw):
            continue
        # remove markers entirely; keep rendered text minus WJ chars
        text = MARKER_RE.sub("", raw).replace(WJ, "")
        _rewrite_paragraph(p, text, [], red=False)
        count += 1
    return count


def main(argv=None):
    ap = argparse.ArgumentParser(description="Zotero-style citations in a .docx.")
    ap.add_argument("--docx", required=True)
    ap.add_argument("--out", help="Output path (default: overwrite input).")
    ap.add_argument("--style", choices=["vancouver", "author-date"], default="vancouver")
    ap.add_argument("--mode", choices=["field", "text"], default="field",
                    help="field: real Zotero fields the Zotero app manages (default); "
                         "text: legacy static text refreshed by this script.")
    ap.add_argument("--action", choices=["refresh", "unlink"], default="refresh")
    ap.add_argument("--heading", default=None,
                    help='Bibliography heading (default "Kaynaklar").')
    ap.add_argument("--no-red", action="store_true",
                    help="Do not color inserted text red (new documents).")
    args = ap.parse_args(argv)

    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    if not os.path.isfile(args.docx):
        print(json.dumps({"error": "not_found", "path": args.docx}, ensure_ascii=False))
        return 0

    doc = docx.Document(args.docx)
    out_path = args.out or args.docx
    red = not args.no_red

    if args.action == "unlink":
        if args.mode == "field":
            print(json.dumps({
                "action": "unlink", "mode": "field",
                "warning": "Canlı Zotero alanlarını dondurma işi Zotero uygulamasınındır "
                           "(Unlink Citations düğmesi). Bu script yalnız text-mode (WJ) "
                           "atıflarını dondurur; alanlara dokunulmadı.",
            }, ensure_ascii=False, indent=2))
        n = unlink(doc)
        doc.save(out_path)
        print(json.dumps({"action": "unlink", "paragraphs_frozen": n, "output": out_path},
                         ensure_ascii=False, indent=2))
        return 0

    lib = load_library()
    heading = args.heading or "Kaynaklar"

    if args.mode == "field":
        account = load_account()
        order, unknown, processed, existing, has_pref, has_bibl = refresh_fields(
            doc, lib, account, args.style, red)
        if not has_pref:
            _insert_pref_field(doc, args.style)
        bib_n = 0
        if not has_bibl:
            new_keys = [k for k in order if k in lib]
            bib_n = write_bibliography_field(doc, lib, new_keys, args.style, red, heading)
        doc.save(out_path)
        print(json.dumps({
            "action": "refresh", "mode": "field", "style": args.style,
            "processed_markers": processed,
            "existing_fields": existing,
            "unique_references": len(order),
            "bibliography_count": bib_n,
            "unknown_keys": unknown,
            "red_revision": red,
            "output": out_path,
            "note": "Alanlar artık Zotero uygulamasının: Word'de Zotero sekmesi → "
                    "Refresh / Document Preferences ile yönetilir.",
        }, ensure_ascii=False, indent=2))
        return 0

    order, unknown, processed = refresh(doc, lib, args.style, red)
    bib_n = write_bibliography(doc, lib, order, args.style, red, heading)
    doc.save(out_path)
    print(json.dumps({
        "action": "refresh", "mode": "text", "style": args.style,
        "processed_markers": processed,
        "unique_references": len(order),
        "bibliography_count": bib_n,
        "unknown_keys": unknown,
        "red_revision": red,
        "output": out_path,
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
