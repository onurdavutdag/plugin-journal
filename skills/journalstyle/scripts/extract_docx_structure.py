#!/usr/bin/env python3
"""
Bir .docx makalenin yapısal özetini çıkarır: başlıklar, tahmini kelime sayısı,
tablo/şekil sayısı, mevcut bölüm başlıkları. journalstyle skill'i bunu
hem uygulama öncesi analiz hem de uygulama sonrası doğrulama için kullanır.

Kullanım:
    python extract_docx_structure.py <dosya.docx>
"""
import sys
import json
from docx import Document


def extract(path):
    doc = Document(path)

    headings = []
    word_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            word_count += len(text.split())
        if para.style.name.startswith("Heading") or para.style.name == "Title":
            headings.append({"style": para.style.name, "text": text})

    table_count = len(doc.tables)

    # Şekil sayısı: inline_shapes docx görsellerini kapsar
    figure_count = len(doc.inline_shapes)

    section_settings = []
    for s in doc.sections:
        section_settings.append({
            "top_margin_cm": round(s.top_margin.cm, 2) if s.top_margin else None,
            "bottom_margin_cm": round(s.bottom_margin.cm, 2) if s.bottom_margin else None,
            "left_margin_cm": round(s.left_margin.cm, 2) if s.left_margin else None,
            "right_margin_cm": round(s.right_margin.cm, 2) if s.right_margin else None,
            "page_width_cm": round(s.page_width.cm, 2) if s.page_width else None,
            "page_height_cm": round(s.page_height.cm, 2) if s.page_height else None,
        })

    result = {
        "word_count_approx": word_count,
        "headings": headings,
        "table_count": table_count,
        "figure_count": figure_count,
        "sections": section_settings,
    }
    return result


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Kullanım: python extract_docx_structure.py <dosya.docx>")
        sys.exit(1)
    result = extract(sys.argv[1])
    print(json.dumps(result, ensure_ascii=False, indent=2))
