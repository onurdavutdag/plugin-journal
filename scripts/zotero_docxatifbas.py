#!/usr/bin/env python3
"""Word citation/bibliography engine — mimics the Zotero Word plugin buttons.

Run by the `journal-s-zotero` agent. Reads the user's real Zotero library via
`zotero_kutuphaneoku.py` (same directory) and processes citation markers in a .docx.

Marker syntax inside the document (either form, one or more keys, ';'-separated):
    {{zref:ITEMKEY}}        {{zref:KEY1;KEY2}}
    [@ITEMKEY]              [@KEY1;KEY2]
ITEMKEY is the 8-char Zotero item key (e.g. F5RI4K5K).

Modes:
  --mode field (default) — every marker becomes a REAL Zotero Word-integration
      field (ADDIN ZOTERO_ITEM CSL_CITATION + ZOTERO_PREF + ZOTERO_BIBL), so the
      user's actual Zotero app recognizes, refreshes and restyles the citations
      from Word (verified against the user's Zotero 7 + Word). Re-running only
      converts NEW markers; existing ZOTERO_* fields belong to the Zotero app
      and are never touched.
  --mode text — legacy static text (WJ-wrapped), refreshed only by this script.

Buttons implemented:
  refresh (default) — Add/Edit Citation + Refresh: replace every marker with the
      in-text citation in the chosen style, (re)number in order of appearance,
      and (re)write the bibliography at the end. Text mode is idempotent within
      the script; field mode hands renumbering to the Zotero app.
  unlink — freeze (text mode): keep the rendered citation text, remove marker
      bookkeeping so the document can no longer be refreshed. For field mode use
      Zotero's own Unlink Citations button instead.

Styles:
  --style vancouver  (default) numeric [1], bibliography in citation order.
      Field mode maps to CSL style id http://www.zotero.org/styles/vancouver.
  --style author-date          (Author, Year), bibliography alphabetical.
      Field mode maps to CSL style id http://www.zotero.org/styles/apa.
  Journal-specific styling beyond these two is handled by the `journal-s-zotero`
  agent (read CSL rules, apply here) — or, in field mode, directly from the Zotero
  app (Document Preferences). Docx citation/bibliography authority lives only in
  that agent; no other skill/agent formats or edits the reference list.

Red-revision rule (global CLAUDE.md): when updating an EXISTING document
(default), all text this script inserts is red (RGB 255,0,0). Pass --no-red for
brand-new documents built from scratch.

Output: JSON report to stdout {processed_markers, unknown_keys, bibliography_count,
output, backup} — exactly ONE JSON object per run.

The source file is never overwritten by default: without --out the result is written
to `<name>_zref.docx` beside it. An explicit --out pointing back at the source takes
a `.bak` copy first. Use the `output` path from the JSON report for the next step.

Formatting is preserved run by run: only the run holding a marker is split, so
italics, bold, super/subscript, hyperlinks and existing Zotero fields survive.

Usage:
    python zotero_docxatifbas.py --docx makale.docx                   # -> makale_zref.docx
    python zotero_docxatifbas.py --docx makale.docx --style author-date
    python zotero_docxatifbas.py --docx makale.docx --mode text --action unlink
    python zotero_docxatifbas.py --docx makale.docx --out makale_atifli.docx --no-red
"""
import argparse
import copy
import json
import os
import random
import re
import shutil
import subprocess
import sys

try:
    import docx  # python-docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    from docx.text.paragraph import Paragraph
except ImportError:
    print(json.dumps({"error": "no_python_docx",
                      "message": "python-docx kurulu değil: pip install python-docx"}))
    sys.exit(0)

RED = RGBColor(0xFF, 0x00, 0x00)
# Case-insensitive: a key typed in lower case is still recognized (it is upper-cased
# on resolution), so a mistyped key surfaces in `unknown_keys` instead of silently
# staying plain text in the document.
MARKER_RE = re.compile(r"\{\{zref:([A-Z0-9;\s]+)\}\}|\[@([A-Z0-9;\s]+)\]", re.IGNORECASE)
BIB_HEADINGS = ("Kaynaklar", "Kaynakça", "References", "Bibliography")
# Invisible bookmark char pair used to keep refresh idempotent: rendered
# citations are wrapped as ⁠{{zref:...}}⁠<visible text>⁠ in a
# single run's text? Too fragile across runs — instead we KEEP the marker in
# the document and render the citation right after it, wrapped in U+2060
# word-joiners, and strip that rendered part on each refresh.
WJ = "⁠"  # word joiner, invisible in Word

_RENDERED_RE = re.compile(WJ + r"[^" + WJ + r"]*" + WJ)


# ------------------------------------------------------------ library -------

def load_library():
    lib_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "zotero_kutuphaneoku.py")
    out = subprocess.run([sys.executable, lib_script, "--items"],
                         capture_output=True, text=True, encoding="utf-8")
    try:
        data = json.loads(out.stdout)
    except json.JSONDecodeError:
        # zotero_kutuphaneoku crashed (traceback on stderr, empty stdout). Honour this script's
        # contract — exactly ONE json object per run — instead of adding a second
        # traceback the calling skill cannot parse.
        raise SystemExit(json.dumps({
            "error": "zotero_kutuphaneoku_failed",
            "message": "zotero_kutuphaneoku.py geçerli JSON döndürmedi; kütüphane okunamadı.",
            "stderr": (out.stderr or "").strip()[-800:],
        }, ensure_ascii=False))
    if isinstance(data, dict) and data.get("error"):
        raise SystemExit(json.dumps(data, ensure_ascii=False))
    return {it["key"]: it for it in data}


def load_account():
    """Account identifiers from zotero_kutuphaneoku --status (for field URIs)."""
    lib_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "zotero_kutuphaneoku.py")
    out = subprocess.run([sys.executable, lib_script, "--status"],
                         capture_output=True, text=True, encoding="utf-8")
    try:
        return json.loads(out.stdout).get("account") or {}
    except (json.JSONDecodeError, AttributeError):
        return {}


# ------------------------------------------------------- Zotero fields ------
# Real Word-integration fields (ADDIN ZOTERO_*) the actual Zotero app manages.
# Structure verified against the user's Zotero 7 + Word (feasibility spike).

STYLE_IDS = {
    "vancouver": "http://www.zotero.org/styles/vancouver",
    "author-date": "http://www.zotero.org/styles/apa",
}

_CSL_TYPES = {
    "journalArticle": "article-journal", "book": "book", "bookSection": "chapter",
    "conferencePaper": "paper-conference", "thesis": "thesis", "report": "report",
    "webpage": "webpage", "preprint": "article", "manuscript": "manuscript",
}

_CITE_INSTR_RE = re.compile(r"ADDIN ZOTERO_ITEM CSL_CITATION\s+(\{.*)", re.S)


def _rand_id(n=8):
    return "".join(random.choice("0123456789ABCDEF") for _ in range(n))


def item_uri(key, account):
    uid = account.get("user_id")
    if uid:
        return f"http://zotero.org/users/{uid}/items/{key}"
    return f"http://zotero.org/users/local/{account.get('local_user_key', 'local')}/items/{key}"


def csl_item_data(it, uri):
    """zotero_kutuphaneoku record -> CSL-JSON itemData (what Zotero embeds in the field)."""
    d = {
        "id": uri,
        "type": _CSL_TYPES.get(it.get("itemType"), "document"),
        "title": it.get("title"),
        "container-title": it.get("container-title"),
        "journalAbbreviation": it.get("journalAbbreviation"),
        "volume": it.get("volume"),
        "issue": it.get("issue"),
        "page": it.get("pages"),
        "DOI": it.get("DOI"),
        "URL": it.get("url"),
    }
    auths = [{"family": c.get("family", ""), "given": c.get("given", "")}
             for c in it.get("creators", []) if c.get("type") in (None, "author")]
    if auths:
        d["author"] = auths
    if it.get("year"):
        d["issued"] = {"date-parts": [[int(it["year"])]]}
    if it.get("PMID"):
        d["PMID"] = it["PMID"]
    return {k: v for k, v in d.items() if v is not None}


def _field_elements(instr, result_text, red, breaks=False, template=None):
    """Elements of one complex Word field (begin|instrText|separate|result|end).

    Returned as a list so the caller decides WHERE they go — appended to the end of
    a paragraph (bibliography, prefs) or inserted exactly where a marker stood.
    `template`: the run whose formatting the visible result inherits.
    """
    els = []

    def fld(t):
        r = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), t)
        r.append(fc)
        els.append(r)

    fld("begin")
    r = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " " + instr + " "
    r.append(it)
    els.append(r)
    fld("separate")
    if result_text:
        lines = result_text.split("\n") if breaks else [result_text]
        els.append(_new_run(lines[0], template, red))
        for line in lines[1:]:
            els.append(_new_run(line, template, red, leading_break=True))
    fld("end")
    return els


def _add_field(p, instr, result_text, red, breaks=False):
    """Append one complex Word field to the end of paragraph p."""
    for el in _field_elements(instr, result_text, red, breaks):
        p._p.append(el)


def _scan_field_instrs(doc):
    """Full instruction text of every complex field, in document order."""
    fields = []
    depth = 0
    for el in doc.element.body.iter():
        if el.tag == qn("w:fldChar"):
            t = el.get(qn("w:fldCharType"))
            if t == "begin":
                depth += 1
                if depth == 1:
                    fields.append("")
            elif t == "end" and depth:
                depth -= 1
        elif el.tag == qn("w:instrText") and depth:
            fields[-1] += el.text or ""
    return fields


def _existing_zotero_state(doc):
    """(citation field count, item keys in order, has_pref, has_bibl)."""
    instrs = _scan_field_instrs(doc)
    keys, n_cites = [], 0
    has_pref = any("ADDIN ZOTERO_PREF" in s for s in instrs)
    has_bibl = any("ADDIN ZOTERO_BIBL" in s for s in instrs)
    for s in instrs:
        m = _CITE_INSTR_RE.search(s)
        if not m:
            continue
        n_cites += 1
        try:
            data = json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            continue
        for ci in data.get("citationItems", []):
            for uri in ci.get("uris", []):
                k = uri.rstrip("/").rsplit("/", 1)[-1].upper()
                if k not in keys:
                    keys.append(k)
    return n_cites, keys, has_pref, has_bibl


def _prefs_xml(style):
    return (
        '<data data-version="3" zotero-version="7.0">'
        f'<session id="{_rand_id(8)}"/>'
        f'<style id="{STYLE_IDS[style]}" locale="en-US" '
        'hasBibliography="1" bibliographyStyleHasBeenSet="0"/>'
        '<prefs><pref name="fieldType" value="Field"/>'
        '<pref name="automaticJournalAbbreviations" value="true"/>'
        '<pref name="noteType" value="0"/></prefs></data>'
    )


def _insert_pref_field(doc, style):
    """Document-preferences field, prepended INTO the document's first paragraph.

    It used to get a paragraph of its own at body index 0. The field carries no
    visible result text, so that paragraph rendered as a blank line above the title —
    field mode is the default, so every output gained one. Putting the field's runs at
    the start of an existing paragraph keeps the same XML plumbing with no extra line.
    Only when the document has no paragraph at all does a new one get created.
    """
    instr = "ADDIN ZOTERO_PREF_1 " + _prefs_xml(style)
    body = doc.element.body
    first = next(iter(body.iter(qn("w:p"))), None)
    if first is None:                       # boş belge: eski davranış
        p = doc.add_paragraph()
        _add_field(p, instr, "", red=False)
        return
    # <w:pPr> must stay the first child of <w:p>, so the runs go right after it.
    pPr = first.find(qn("w:pPr"))
    at = list(first).index(pPr) + 1 if pPr is not None else 0
    for offset, el in enumerate(_field_elements(instr, "", red=False)):
        first.insert(at + offset, el)


def citation_field_instr(keys, lib, account, style, order):
    """Build the ADDIN ZOTERO_ITEM instruction + visible result for one marker."""
    citation_items, nums_or_cites = [], []
    for k in keys:
        it = lib[k]
        uri = item_uri(k, account)
        citation_items.append({"id": uri, "uris": [uri], "itemData": csl_item_data(it, uri)})
        if style == "vancouver":
            nums_or_cites.append(str(order.index(k) + 1))
        else:
            nums_or_cites.append(author_date_intext(it)[1:-1])
    if style == "vancouver":
        visible = "[" + ",".join(nums_or_cites) + "]"
    else:
        visible = "(" + "; ".join(nums_or_cites) + ")"
    citation = {
        "citationID": _rand_id(8),
        "properties": {"formattedCitation": visible, "plainCitation": visible, "noteIndex": 0},
        "citationItems": citation_items,
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    return "ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(citation, ensure_ascii=False), visible


def refresh_fields(doc, lib, account, style, red):
    """Convert every marker into a real Zotero citation field.

    Existing ADDIN ZOTERO_* fields are the Zotero app's property — never touched.
    Returns (order, unknown, processed, existing_cites).
    """
    existing_cites, order, has_pref, has_bibl = _existing_zotero_state(doc)
    unknown, processed = [], 0

    for p in _iter_paragraphs(doc):
        raw = _para_text(p)
        if not MARKER_RE.search(raw) and WJ not in raw:
            continue

        # legacy text-mode renderings are dropped; everything else stays untouched
        edits = [(m.start(), m.end(), None) for m in _RENDERED_RE.finditer(raw)]

        for m in MARKER_RE.finditer(raw):
            keys = [k.strip().upper() for k in (m.group(1) or m.group(2)).split(";") if k.strip()]
            good = []
            for k in keys:
                if k in lib:
                    good.append(k)
                    if k not in order:
                        order.append(k)
                elif k not in unknown:
                    unknown.append(k)
            if not good:
                continue  # unknown key(s): the marker is left in place, nothing lost
            instr, visible = citation_field_instr(good, lib, account, style, order)
            edits.append((m.start(), m.end(),
                          lambda tpl, i=instr, v=visible: _field_elements(i, v, red, template=tpl)))
            processed += 1

        _apply_edits(p, edits)

    return order, unknown, processed, existing_cites, has_pref, has_bibl


def write_bibliography_field(doc, lib, order, style, red, heading):
    """Heading paragraph + one ADDIN ZOTERO_BIBL field with rendered entries."""
    if not order:
        return 0
    h = doc.add_paragraph()
    run = h.add_run(heading)
    run.bold = True
    if red:
        run.font.color.rgb = RED
    items = [lib[k] for k in order if k in lib]
    if style == "author-date":
        items = sorted(items, key=_sort_key_author_date)
        entries = [author_date_entry(it) for it in items]
    else:
        entries = [f"{i}. {vancouver_entry(it)}" for i, it in enumerate(items, 1)]
    p = doc.add_paragraph()
    _add_field(p, 'ADDIN ZOTERO_BIBL {"uncited":[],"omitted":[],"custom":[]} CSL_BIBLIOGRAPHY',
               "\n".join(entries), red, breaks=True)
    return len(entries)


# ------------------------------------------------------------ formatting ----

def _authors_vancouver(creators):
    """Surname Initials, first 6 then et al. (citation-format.md rule)."""
    auths = [c for c in creators if c.get("type") == "author"] or creators
    names = []
    for c in auths[:6]:
        fam = c.get("family", "").strip()
        giv = c.get("given", "").strip()
        initials = "".join(w[0] for w in re.split(r"[\s\-.]+", giv) if w)
        names.append((fam + " " + initials).strip())
    s = ", ".join(n for n in names if n)
    if len(auths) > 6:
        s += ", et al."
    return s


def vancouver_entry(it):
    """One numbered reference-list entry (without the leading number)."""
    parts = []
    a = _authors_vancouver(it.get("creators", []))
    if a:
        parts.append(a + ".")
    if it.get("title"):
        parts.append(it["title"].rstrip(".") + ".")
    j = it.get("journalAbbreviation") or it.get("container-title")
    if j:
        parts.append(j + ".")
    tail = ""
    if it.get("year"):
        tail = it["year"]
    if it.get("volume"):
        tail += ";" + it["volume"]
        if it.get("issue"):
            tail += "(" + it["issue"] + ")"
    if it.get("pages"):
        tail += ":" + it["pages"]
    if tail:
        parts.append(tail + ".")
    if it.get("DOI"):
        parts.append("doi:" + it["DOI"] + ".")
    if it.get("PMID"):
        parts.append("PMID: " + it["PMID"] + ".")
    return " ".join(parts)


def author_date_intext(it):
    auths = [c for c in it.get("creators", []) if c.get("type") == "author"] or it.get("creators", [])
    if not auths:
        fam = it.get("title", "Anon")[:20]
    elif len(auths) == 1:
        fam = auths[0]["family"]
    elif len(auths) == 2:
        fam = auths[0]["family"] + " & " + auths[1]["family"]
    else:
        fam = auths[0]["family"] + " et al."
    return "(" + fam + ", " + (it.get("year") or "t.y.") + ")"


def author_date_entry(it):
    """APA-flavoured author-date reference entry."""
    auths = [c for c in it.get("creators", []) if c.get("type") == "author"] or it.get("creators", [])
    names = []
    for c in auths:
        giv = c.get("given", "").strip()
        initials = ". ".join(w[0] for w in re.split(r"[\s\-.]+", giv) if w)
        names.append(c.get("family", "") + (", " + initials + "." if initials else ""))
    astr = ", ".join(names[:-1]) + (", & " + names[-1] if len(names) > 1 else (names[0] if names else ""))
    parts = [astr, "(" + (it.get("year") or "t.y.") + ")."]
    if it.get("title"):
        parts.append(it["title"].rstrip(".") + ".")
    j = it.get("container-title") or it.get("journalAbbreviation")
    if j:
        seg = j
        if it.get("volume"):
            seg += ", " + it["volume"]
            if it.get("issue"):
                seg += "(" + it["issue"] + ")"
        if it.get("pages"):
            seg += ", " + it["pages"]
        parts.append(seg + ".")
    if it.get("DOI"):
        parts.append("https://doi.org/" + it["DOI"])
    return " ".join(p for p in parts if p)


def _sort_key_author_date(it):
    cs = it.get("creators", [])
    fam = cs[0]["family"].lower() if cs else ""
    return (fam, it.get("year") or "")


# ------------------------------------------------------------ docx ops ------

def _iter_paragraphs(doc):
    """Every paragraph in TRUE DOCUMENT ORDER — body and table cells interleaved.

    Walking `doc.paragraphs` first and `doc.tables` after put every table citation
    behind every body citation, so a table in the middle of the manuscript got a
    higher Vancouver number than the text following it. `body.iter()` visits the XML
    in document order, and a `<w:p>` inside a `<w:tbl>` therefore arrives in its real
    position.
    """
    body = doc.element.body
    for el in body.iter(qn("w:p")):
        yield Paragraph(el, doc)


def _run_elements(p):
    """Every <w:r> of the paragraph in document order — INCLUDING the runs inside
    a <w:hyperlink>, which `Paragraph.runs` (direct children only) does not see."""
    return p._element.xpath(".//w:r")


def _run_text(r_el):
    """Visible text of a run element (same rules python-docx uses for Run.text).

    Field instructions (<w:instrText>) and field characters contribute nothing, so
    an existing Zotero field's plumbing has zero length and is never split/removed.
    """
    out = []
    for child in r_el:
        tag = child.tag
        if tag == qn("w:t"):
            out.append(child.text or "")
        elif tag == qn("w:tab"):
            out.append("\t")
        elif tag in (qn("w:br"), qn("w:cr")):
            out.append("\n")
    return "".join(out)


def _para_text(p):
    return "".join(_run_text(r) for r in _run_elements(p))


def _run_offsets(p):
    """[(run element, start, end)] over the paragraph's concatenated text."""
    offs, pos = [], 0
    for r in _run_elements(p):
        n = len(_run_text(r))
        offs.append((r, pos, pos + n))
        pos += n
    return offs


def _child_len(el):
    if el.tag == qn("w:t"):
        return len(el.text or "")
    if el.tag in (qn("w:tab"), qn("w:br"), qn("w:cr")):
        return 1
    return 0


def _new_run(text, template=None, red=False, leading_break=False):
    """A run carrying `text` with `template`'s formatting (rPr) inherited."""
    r = OxmlElement("w:r")
    if template is not None:
        rPr = template.find(qn("w:rPr"))
        if rPr is not None:
            r.append(copy.deepcopy(rPr))
    if leading_break:
        r.append(OxmlElement("w:br"))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    if red:
        _set_red(r)
    return r


def _set_red(r_el):
    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_el.insert(0, rPr)
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), "FF0000")


def _split_run(r_el, k):
    """Split r_el after k characters, keeping every formatting property on both
    halves. The right-hand half is inserted straight after the left one (inside
    the same parent, so a run inside a hyperlink stays inside it)."""
    if k <= 0:
        return None
    right = OxmlElement("w:r")
    rPr = r_el.find(qn("w:rPr"))
    if rPr is not None:
        right.append(copy.deepcopy(rPr))
    pos = 0
    for child in list(r_el):
        if child.tag == qn("w:rPr"):
            continue
        n = _child_len(child)
        if pos >= k:                       # entirely after the cut
            r_el.remove(child)
            right.append(child)
        elif pos + n <= k:                 # entirely before the cut
            pass
        elif child.tag == qn("w:t"):       # straddles the cut (only <w:t> can)
            txt = child.text or ""
            cut = k - pos
            child.text = txt[:cut]
            child.set(qn("xml:space"), "preserve")
            nt = OxmlElement("w:t")
            nt.set(qn("xml:space"), "preserve")
            nt.text = txt[cut:]
            right.append(nt)
        pos += n
    if all(c.tag == qn("w:rPr") for c in right):   # nothing landed on the right
        return None
    r_el.addnext(right)
    return right


def _replace_span(p, start, end, builder):
    """Replace the paragraph text range [start, end) with the elements returned by
    `builder(template_run)` — or delete it when builder is None.

    Only the runs the range actually touches are split/removed; every other run
    (italics, bold, super/subscript, hyperlinks, existing Zotero fields) is left
    exactly where it was.
    """
    for r, s, e in _run_offsets(p):        # split at the right edge first …
        if s < end < e:
            _split_run(r, end - s)
            break
    for r, s, e in _run_offsets(p):        # … then at the left edge
        if s < start < e:
            _split_run(r, start - s)
            break
    covered = [r for r, s, e in _run_offsets(p) if e > s and s >= start and e <= end]
    if not covered:
        return False
    anchor = covered[0]
    for el in (builder(anchor) if builder else []):
        anchor.addprevious(el)
    for r in covered:
        r.getparent().remove(r)
    return True


def _apply_edits(p, edits):
    """Apply (start, end, builder) edits right-to-left so offsets stay valid."""
    applied = 0
    for start, end, builder in sorted(edits, key=lambda e: e[0], reverse=True):
        if _replace_span(p, start, end, builder):
            applied += 1
    return applied


def _is_ours(p):
    """A paragraph this script wrote carries an invisible word-joiner tag."""
    return WJ in _para_text(p)


def _remove_old_bibliography(doc):
    """Delete the bibliography section previously written by this script.

    Bounded on both ends: it starts at a WJ-tagged bibliography heading and stops
    at the first paragraph that is NOT WJ-tagged, so anything the user keeps after
    the reference list (tables, figure captions, appendix, acknowledgements) is
    never touched.
    """
    body_paras = doc.paragraphs
    start = None
    for i, p in enumerate(body_paras):
        if p.text.strip().strip(WJ) in BIB_HEADINGS and WJ in p.text:
            start = i
            break
    if start is None:
        return 0
    removed = 0
    for p in body_paras[start:]:
        if removed and not _is_ours(p):
            break
        el = p._element
        el.getparent().remove(el)
        removed += 1
    return removed


# ------------------------------------------------------------ main flow -----

def _resolve_out_path(docx_path, out):
    """Default output: NEVER the source file — `<name>_zref.docx` beside it."""
    if out:
        return out
    stem, ext = os.path.splitext(docx_path)
    return stem + "_zref" + ext


def _backup_if_overwriting(src, out):
    """An explicit --out pointing back at the source still gets a .bak first."""
    if os.path.isfile(src) and os.path.abspath(src) == os.path.abspath(out):
        bak = src + ".bak"
        shutil.copy2(src, bak)
        return bak
    return None


def refresh(doc, lib, style, red):
    """Render/refresh all markers; return (order, unknown, count)."""
    order = []          # item keys in order of first appearance
    unknown = []
    processed = 0

    for p in _iter_paragraphs(doc):
        raw = _para_text(p)
        if WJ not in raw and not MARKER_RE.search(raw):
            continue

        # previous renderings are dropped and re-made (idempotent refresh)
        edits = [(m.start(), m.end(), None) for m in _RENDERED_RE.finditer(raw)]

        for m in MARKER_RE.finditer(raw):
            keys = [k.strip().upper() for k in (m.group(1) or m.group(2)).split(";") if k.strip()]
            nums_or_cites = []
            for k in keys:
                it = lib.get(k)
                if it is None:
                    if k not in unknown:
                        unknown.append(k)
                    continue
                if k not in order:
                    order.append(k)
                if style == "vancouver":
                    nums_or_cites.append(str(order.index(k) + 1))
                else:
                    # bare "Author, Year" — merged into one paren below
                    nums_or_cites.append(author_date_intext(it)[1:-1])
            if not nums_or_cites:
                # No key resolved: leave the marker exactly where it is instead of
                # rendering "[?]" into the manuscript. Field mode already behaved this
                # way; the key still surfaces in `unknown_keys`, so nothing is hidden.
                continue
            if style == "vancouver":
                cite = "[" + ",".join(nums_or_cites) + "]"
            else:
                cite = "(" + "; ".join(nums_or_cites) + ")"
            marker = raw[m.start():m.end()]

            # marker stays (refresh needs it); only the visible citation is red
            def _render(tpl, mk=marker, ct=cite):
                return [_new_run(mk + WJ, tpl),
                        _new_run(ct, tpl, red=red),
                        _new_run(WJ, tpl)]

            edits.append((m.start(), m.end(), _render))
            processed += 1

        _apply_edits(p, edits)

    return order, unknown, processed


def write_bibliography(doc, lib, order, style, red, heading):
    if not order:
        # nothing to write: never delete an existing reference list on an empty run
        return 0
    _remove_old_bibliography(doc)
    h = doc.add_paragraph()
    run = h.add_run(heading + WJ)  # WJ tags it as ours for future removal
    run.bold = True
    if red:
        run.font.color.rgb = RED
    items = [lib[k] for k in order]
    if style == "author-date":
        items = sorted(items, key=_sort_key_author_date)
        entries = [author_date_entry(it) for it in items]
    else:
        entries = [f"{i}. {vancouver_entry(it)}" for i, it in enumerate(items, 1)]
    for e in entries:
        p = doc.add_paragraph()
        run = p.add_run(e + WJ)  # every entry tagged too, so removal stays bounded
        if red:
            run.font.color.rgb = RED
    return len(entries)


def unlink(doc):
    """Freeze rendered citations: drop markers + WJ wrappers, keep visible text."""
    count = 0
    for p in _iter_paragraphs(doc):
        raw = _para_text(p)
        if WJ not in raw and not MARKER_RE.search(raw):
            continue
        edits, covered = [], set()
        for m in MARKER_RE.finditer(raw):                 # markers go entirely
            edits.append((m.start(), m.end(), None))
            covered.update(range(m.start(), m.end()))
        for m in _RENDERED_RE.finditer(raw):              # keep the text, drop the WJ pair
            inner = raw[m.start() + 1:m.end() - 1]
            edits.append((m.start(), m.end(),
                          (lambda tpl, t=inner: [_new_run(t, tpl)]) if inner else None))
            covered.update(range(m.start(), m.end()))
        for i, ch in enumerate(raw):                      # any stray, unpaired WJ
            if ch == WJ and i not in covered:
                edits.append((i, i + 1, None))
        if not edits:
            continue
        _apply_edits(p, edits)
        count += 1
    return count


def main(argv=None):
    ap = argparse.ArgumentParser(description="Zotero-style citations in a .docx.")
    ap.add_argument("--docx", required=True)
    ap.add_argument("--out", help="Output path (default: <ad>_zref.docx beside the input; "
                                  "the source is never overwritten silently).")
    ap.add_argument("--style", choices=["vancouver", "author-date"], default="vancouver")
    ap.add_argument("--mode", choices=["field", "text"], default="field",
                    help="field: real Zotero fields the Zotero app manages (default); "
                         "text: legacy static text refreshed by this script.")
    ap.add_argument("--action", choices=["refresh", "unlink"], default="refresh")
    ap.add_argument("--heading", default=None,
                    help='Bibliography heading (default "Kaynaklar").')
    ap.add_argument("--no-red", action="store_true",
                    help="Do not color inserted text red (new documents).")
    args = ap.parse_args(argv)

    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    if not os.path.isfile(args.docx):
        print(json.dumps({"error": "not_found", "path": args.docx}, ensure_ascii=False))
        return 0

    doc = docx.Document(args.docx)
    out_path = _resolve_out_path(args.docx, args.out)
    red = not args.no_red

    if args.action == "unlink":
        if args.mode == "field":
            # single JSON on stdout, and the document is left untouched
            print(json.dumps({
                "action": "unlink", "mode": "field",
                "paragraphs_frozen": 0, "output": None,
                "warning": "Canlı Zotero alanlarını dondurma işi Zotero uygulamasınındır "
                           "(Unlink Citations düğmesi). Bu script yalnız text-mode (WJ) "
                           "atıflarını dondurur; alanlara dokunulmadı, dosya kaydedilmedi. "
                           "Text-mode atıfları dondurmak için: --mode text --action unlink.",
            }, ensure_ascii=False, indent=2))
            return 0
        n = unlink(doc)
        backup = _backup_if_overwriting(args.docx, out_path)
        doc.save(out_path)
        print(json.dumps({"action": "unlink", "mode": "text", "paragraphs_frozen": n,
                          "output": out_path, "backup": backup},
                         ensure_ascii=False, indent=2))
        return 0

    lib = load_library()
    heading = args.heading or "Kaynaklar"

    if args.mode == "field":
        account = load_account()
        order, unknown, processed, existing, has_pref, has_bibl = refresh_fields(
            doc, lib, account, args.style, red)
        if not has_pref:
            _insert_pref_field(doc, args.style)
        bib_n = 0
        if not has_bibl:
            new_keys = [k for k in order if k in lib]
            bib_n = write_bibliography_field(doc, lib, new_keys, args.style, red, heading)
        backup = _backup_if_overwriting(args.docx, out_path)
        doc.save(out_path)
        print(json.dumps({
            "action": "refresh", "mode": "field", "style": args.style,
            "processed_markers": processed,
            "existing_fields": existing,
            "unique_references": len(order),
            "bibliography_count": bib_n,
            "unknown_keys": unknown,
            "red_revision": red,
            "output": out_path,
            "backup": backup,
            "note": "Alanlar artık Zotero uygulamasının: Word'de Zotero sekmesi → "
                    "Refresh / Document Preferences ile yönetilir.",
        }, ensure_ascii=False, indent=2))
        return 0

    order, unknown, processed = refresh(doc, lib, args.style, red)
    bib_n = write_bibliography(doc, lib, order, args.style, red, heading)
    backup = _backup_if_overwriting(args.docx, out_path)
    doc.save(out_path)
    print(json.dumps({
        "action": "refresh", "mode": "text", "style": args.style,
        "processed_markers": processed,
        "unique_references": len(order),
        "bibliography_count": bib_n,
        "unknown_keys": unknown,
        "red_revision": red,
        "output": out_path,
        "backup": backup,
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
